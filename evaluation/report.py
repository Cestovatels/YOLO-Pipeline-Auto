import os
import numpy as np
import pandas as pd
from datetime import datetime
from string import Template
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

text_template = Template(
"""
\n- $image_count Radiographies
- $label_count Labels
- %10 Test - %10 Validation - %80 Training
- Trained $task model for $epoch epochs with $model and $imgsz image size

mAP_0.5: $map50
mAP_0.5:0.95: $map5095

True Positive: $tp
False Positive: $fp
False Negative: $fn
Overall Accuracy: $accuracy

Jaccards, Dice: 
$jaccards_dice
Overall Dice: $overall_dice
Overall Jaccard: $overall_jaccard

IoU: 
$IoU
overall_IoU: $overall_IoU


Test image count: $test_image_count
Train image count: $train_image_count
Validation image count: $valid_image_count

Test label count: $test_label_count
Train label count: $train_label_count
Validation label count: $valid_label_count
""")

def create_word_file(metric_dict, training_name, task, selected_model, epochs, imgsz,
                     map50, map5095, jaccards_dice, overall_dice, overall_jaccard,
                     IoU, overall_IoU):
    
    # augments = '\n'.join([str(x) for x in custom_transforms])
    
    train_image_count = len(os.listdir(f"Results-Yolo-Auto/{training_name}/Dataset_{training_name}/train/images"))
    valid_image_count = len(os.listdir(f"Results-Yolo-Auto/{training_name}/Dataset_{training_name}/valid/images"))
    test_image_count = len(os.listdir(f"Results-Yolo-Auto/{training_name}/Dataset_{training_name}/test/images"))

    train_label_count = sum(d["train_label_count"] for d in metric_dict.values())
    valid_label_count = sum(d["valid_label_count"] for d in metric_dict.values())
    test_label_count = sum(d["test_label_count"] for d in metric_dict.values())

    true_positive = int(sum(d["tp"] for d in metric_dict.values()))
    false_positive = int(sum(d["fp"] for d in metric_dict.values()))
    false_negative = int(sum(d["fn"] for d in metric_dict.values()))
    
    accuracy = true_positive / (true_positive + false_positive + false_negative + 1e-9)

    doc = Document()
    
    # Title
    p = doc.add_paragraph()
    run = p.add_run(training_name)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Main content
    formatted_text = text_template.substitute(
        image_count=train_image_count + valid_image_count + test_image_count,
        label_count=train_label_count + valid_label_count + test_label_count,
        task=task,
        model=selected_model,
        epoch=epochs,
        imgsz=imgsz,
        # augment=augments,
        test_image_count=test_image_count,
        train_image_count=train_image_count,
        valid_image_count=valid_image_count,
        test_label_count=test_label_count,
        train_label_count=train_label_count,
        valid_label_count=valid_label_count,
        map50=map50,
        map5095=map5095,
        tp=true_positive,
        fp=false_positive,
        fn=false_negative,
        accuracy=accuracy,
        jaccards_dice=jaccards_dice,
        overall_dice=overall_dice,
        overall_jaccard=overall_jaccard,
        IoU=IoU,
        overall_IoU=overall_IoU
    )

    doc.add_paragraph(formatted_text)
    doc.save(f"Results-Yolo-Auto/{training_name}/report.docx")

def create_report(metrics, task, training_name, selected_model, epochs, imgsz, jaccards_dice,
                  overall_dice, overall_jaccard, IoU, overall_IoU):
    
    def calculate_label_counts(metric_dict):
        for folder in ["train", "valid", "test"]:
            label_dir = f"Results-Yolo-Auto/{training_name}/Dataset_{training_name}/{folder}/labels"
            labels = os.listdir(label_dir)
            for k in metric_dict:
                metric_dict[k][f"{folder}_image_count"] = len(labels)
            for label in labels:
                with open(os.path.join(label_dir, label), "r") as f:
                    for line in f:
                        class_id = int(line.split()[0])
                        class_name = metrics.names[class_id]
                        metric_dict[class_name][f"{folder}_label_count"] += 1
        return metric_dict

    metric_dict = {
        name: {
            "train_label_count": 0,
            "valid_label_count": 0,
            "test_label_count": 0,
            "algorithm": selected_model,
            "learning_rate": 0.01,
            "epochs": epochs,
            "image_size": imgsz
        } for name in metrics.names.values()
    }

    metric_dict = calculate_label_counts(metric_dict)

    for i, name in enumerate(metrics.names.values()):
        tp = metrics.confusion_matrix.matrix[i][i]
        fp = np.sum(metrics.confusion_matrix.matrix[i, :]) - tp
        fn = np.sum(metrics.confusion_matrix.matrix[:, i]) - tp

        precision = tp / (tp + fp + 1e-9)
        recall = tp / (tp + fn + 1e-9)
        f1 = 2 * precision * recall / (precision + recall + 1e-9)
        dice = 2 * tp / (2 * tp + fp + fn + 1e-9)

        metric_dict[name].update({
            "tp": tp,
            "fp": fp,
            "fn": fn,
            "precision": precision,
            "recall": recall,
            "f1-score": f1,
            "dice-score": dice
        })

    df = pd.DataFrame.from_dict(metric_dict, orient='index')
    df.to_excel(f"Results-Yolo-Auto/{training_name}/report_excel.xlsx", index=True)

    map50 = metrics.results_dict.get('metrics/mAP50(M)', metrics.results_dict.get('metrics/mAP50(B)', 0.0))
    map5095 = metrics.results_dict.get('metrics/mAP50-95(M)', metrics.results_dict.get('metrics/mAP50-95(B)', 0.0))

    create_word_file(
        metric_dict, training_name, task, selected_model, epochs, imgsz,
        map50, map5095, jaccards_dice, overall_dice, overall_jaccard,
        IoU, overall_IoU
    )



